"""Generate a test .docx that exercises every link shape stage 1 must handle."""
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single")
    rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = text
    r.append(t)
    h.append(r)
    paragraph._p.append(h)


def add_fld_simple_link(paragraph, url, text):
    f = OxmlElement("w:fldSimple")
    f.set(qn("w:instr"), f' HYPERLINK "{url}" ')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t"); t.text = text
    r.append(t)
    f.append(r)
    paragraph._p.append(f)


def add_complex_field_link(paragraph, url, text):
    def fld(kind):
        r = OxmlElement("w:r")
        fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), kind)
        r.append(fc)
        return r
    paragraph._p.append(fld("begin"))
    r = OxmlElement("w:r")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f' HYPERLINK "{url}" '
    r.append(it)
    paragraph._p.append(r)
    paragraph._p.append(fld("separate"))
    r = OxmlElement("w:r")
    t = OxmlElement("w:t"); t.text = text
    r.append(t)
    paragraph._p.append(r)
    paragraph._p.append(fld("end"))


doc = docx.Document()
doc.add_heading("Extraction test document", level=1)

p = doc.add_paragraph("Claude is a very good AI and can do lots of things (")
add_hyperlink(p, "https://www.anthropic.com/claude", "Claude")
p.add_run("). Government reports are where translators struggle most (")
add_hyperlink(
    p,
    "https://www.health.gov.au/resources/publications/example-report",
    "Department of Health 2025",
)
p.add_run(").")

p = doc.add_paragraph("An older-style simple field link: (")
add_fld_simple_link(p, "https://www.abs.gov.au/statistics/example", "ABS 2024")
p.add_run(") and a complex field link: (")
add_complex_field_link(p, "https://doi.org/10.1000/example123", "Smith 2023")
p.add_run(").")

# duplicate URL in a later paragraph — dedup fodder for stage 5
p = doc.add_paragraph("The same source cited again later (")
add_hyperlink(p, "https://www.anthropic.com/claude", "Claude, above")
p.add_run(").")

# a link inside a table cell
table = doc.add_table(rows=1, cols=2)
cell_p = table.cell(0, 1).paragraphs[0]
cell_p.add_run("Source: ")
add_hyperlink(cell_p, "https://www.aihw.gov.au/reports/example", "AIHW report")
table.cell(0, 0).paragraphs[0].add_run("Table citation")

# non-web + internal links, which must be skipped/counted, not extracted
p = doc.add_paragraph("Contact (")
add_hyperlink(p, "mailto:someone@example.gov.au", "email us")
p.add_run(") — not a citation.")

doc.save("dev/test_document.docx")
print("wrote dev/test_document.docx")
